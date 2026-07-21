import pytest
import io
import docx
from parsers.transcript_parser import parse_transcript

def test_normal_transcript_txt():
    """Verify parser behavior with a normal speaker-labeled text transcript."""
    transcript_content = (
        "John Doe: Welcome to the project sync.\n"
        "Alice Smith: Thanks, John. Glad to be here.\n"
        "John Doe (10:02): Let's start with the upload module discussion.\n"
        "Alice Smith: Yes, I completed the endpoint implementation yesterday."
    ).encode("utf-8")
    
    result = parse_transcript(transcript_content, ".txt")
    
    assert result["has_speaker_labels"] is True
    assert result["speakers"] == ["Alice Smith", "John Doe"]
    assert len(result["chunks"]) == 1
    assert "John Doe: Welcome to the project sync." in result["raw_text"]
    assert "Alice Smith: Yes, I completed the endpoint implementation yesterday." in result["raw_text"]

def test_normal_transcript_docx():
    """Verify parser behavior with a normal speaker-labeled Word (.docx) transcript."""
    # Generate a temporary docx in-memory
    doc = docx.Document()
    doc.add_paragraph("John: Hello everyone.")
    doc.add_paragraph("Alice (10:01): Hi John, how are you?")
    doc.add_paragraph("John: I am doing great.")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    docx_bytes = file_stream.getvalue()
    
    result = parse_transcript(docx_bytes, ".docx")
    
    assert result["has_speaker_labels"] is True
    assert result["speakers"] == ["Alice", "John"]
    assert len(result["chunks"]) == 1
    assert "John: Hello everyone." in result["raw_text"]
    assert "Alice: Hi John, how are you?" in result["raw_text"]

def test_no_speaker_labels():
    """Verify that a transcript without speaker labels is handled gracefully."""
    transcript_content = (
        "Welcome to the meeting.\n"
        "We are here to discuss project updates.\n"
        "Please review your action items before Friday."
    ).encode("utf-8")
    
    result = parse_transcript(transcript_content, ".txt")
    
    assert result["has_speaker_labels"] is False
    assert result["speakers"] == []
    assert len(result["chunks"]) == 1
    assert "Welcome to the meeting." in result["raw_text"]

def test_empty_transcript():
    """Verify that empty inputs (or whitespace only) raise ValueError."""
    # Test completely empty bytes
    with pytest.raises(ValueError, match="Empty transcript"):
        parse_transcript(b"", ".txt")
        
    # Test whitespace bytes
    with pytest.raises(ValueError, match="Empty transcript"):
        parse_transcript(b"   \n  \n  ", ".txt")

def test_unsupported_file_extension():
    """Verify that unsupported extensions raise ValueError."""
    with pytest.raises(ValueError, match="Unsupported file format"):
        parse_transcript(b"some content", ".pdf")

def test_very_long_transcript_chunking():
    """Verify chunking behavior with a long transcript and small max_chunk_chars."""
    # We will pass a transcript and set max_chunk_chars to 50
    transcript_content = (
        "Alice: Line number one is here.\n"   # len: 31 chars (+1 newline) = 32
        "Bob: Line number two is longer.\n"   # len: 31 chars (+1 newline) = 32
        "Alice: Line number three is long.\n" # len: 33 chars (+1 newline) = 34
        "Bob: And line number four is here."  # len: 34 chars = 34
    ).encode("utf-8")
    
    # Let's call with max_chunk_chars=50.
    # Chunk 1: "Alice: Line number one is here." (len 31)
    # Next line "Bob: Line number two is longer." (len 31) -> 31 + 1 + 31 = 63 > 50 -> goes to Chunk 2.
    # Chunk 2: "Bob: Line number two is longer." (len 31)
    # Next line "Alice: Line number three is long." (len 33) -> 31 + 1 + 33 = 65 > 50 -> goes to Chunk 3.
    # Chunk 3: "Alice: Line number three is long." (len 33)
    # Next line "Bob: And line number four is here." (len 34) -> 33 + 1 + 34 = 68 > 50 -> goes to Chunk 4.
    # Chunk 4: "Bob: And line number four is here." (len 34)
    # So we should get 4 chunks.
    result = parse_transcript(transcript_content, ".txt", max_chunk_chars=50)
    
    assert len(result["chunks"]) == 4
    assert result["chunks"][0] == "Alice: Line number one is here."
    assert result["chunks"][1] == "Bob: Line number two is longer."
    assert result["chunks"][2] == "Alice: Line number three is long."
    assert result["chunks"][3] == "Bob: And line number four is here."
    
    # If we call with max_chunk_chars=70
    # Chunk 1: "Alice: Line number one is here." + "\n" + "Bob: Line number two is longer." (len 31 + 1 + 31 = 63 <= 70)
    # Chunk 2: "Alice: Line number three is long." + "\n" + "Bob: And line number four is here." (len 33 + 1 + 34 = 68 <= 70)
    # So we should get 2 chunks.
    result2 = parse_transcript(transcript_content, ".txt", max_chunk_chars=70)
    assert len(result2["chunks"]) == 2
    assert result2["chunks"][0] == "Alice: Line number one is here.\nBob: Line number two is longer."
    assert result2["chunks"][1] == "Alice: Line number three is long.\nBob: And line number four is here."
